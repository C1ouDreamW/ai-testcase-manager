from io import BytesIO

from docx import Document

MAX_FILE_SIZE = 10 * 1024 * 1024
ALLOWED_EXTENSIONS = {".docx", ".md", ".markdown"}


class DocumentParseError(ValueError):
    """文档解析失败时抛出的异常。"""


def _decode_text(data: bytes) -> str:
    """将字节数据按多种编码尝试解码为字符串。

    依次尝试 UTF-8-BOM、UTF-8、GBK 编码，全部失败则抛出异常。

    Args:
        data (bytes): 要解码的字节数据。

    Returns:
        str: 解码后的字符串。

    Raises:
        DocumentParseError: 所有编码均无法识别时抛出。
    """
    for encoding in ("utf-8-sig", "utf-8", "gbk"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentParseError("无法识别 Markdown 文件编码，请使用 UTF-8")


def parse_docx(data: bytes) -> str:
    """解析 .docx 文件，提取段落文本和表格内容。

    Args:
        data (bytes): .docx 文件的原始字节数据。

    Returns:
        str: 提取的文本内容，段落之间以双换行分隔，表格行以竖线分隔。

    Raises:
        DocumentParseError: 文件格式错误或未提取到文本时抛出。
    """
    try:
        doc = Document(BytesIO(data))
    except Exception as exc:
        raise DocumentParseError("Word 文件解析失败，请确认是 .docx 格式") from exc

    parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    content = "\n\n".join(parts).strip()
    if not content:
        raise DocumentParseError("Word 文件中未提取到文本内容")
    return content


def parse_markdown(data: bytes) -> str:
    """解析 Markdown 文件，自动识别编码并解码为文本。

    Args:
        data (bytes): Markdown 文件的原始字节数据。

    Returns:
        str: 解码后的文本内容。

    Raises:
        DocumentParseError: 编码识别失败或文件内容为空时抛出。
    """
    content = _decode_text(data).strip()
    if not content:
        raise DocumentParseError("Markdown 文件内容为空")
    return content


def parse_upload(filename: str, data: bytes) -> tuple[str, str]:
    """根据文件扩展名分发到对应的解析器，返回文本内容和文件类型。

    校验文件大小和扩展名，支持 .docx、.md、.markdown 格式。

    Args:
        filename (str): 上传文件的原始文件名。
        data (bytes): 文件的原始字节数据。

    Returns:
        tuple[str, str]: (解析后的文本内容, 文件类型标识)。

    Raises:
        DocumentParseError: 文件过大或格式不支持时抛出。
    """
    if len(data) > MAX_FILE_SIZE:
        raise DocumentParseError("文件大小不能超过 10MB")

    name = filename.lower()
    if name.endswith(".docx"):
        return parse_docx(data), "docx"
    if name.endswith(".md") or name.endswith(".markdown"):
        return parse_markdown(data), "markdown"

    if name.endswith(".doc"):
        raise DocumentParseError("暂不支持 .doc 格式，请另存为 .docx 后上传")

    raise DocumentParseError("仅支持 .docx、.md、.markdown 格式")


def title_from_filename(filename: str) -> str:
    """从文件名提取标题，去除扩展名。

    Args:
        filename (str): 文件名。

    Returns:
        str: 去除扩展名后的标题文本。
    """
    if "." in filename:
        return filename.rsplit(".", 1)[0]
    return filename
